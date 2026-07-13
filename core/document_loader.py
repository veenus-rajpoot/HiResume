"""
Turns whatever the user uploads (pdf / docx / txt / pasted text) into
plain text that can be chunked and embedded.
"""
from __future__ import annotations

import io
from pypdf import PdfReader
from docx import Document as DocxDocument


def load_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def load_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # also pull table cell text (many people keep project lists in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def load_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch on file extension. Raises ValueError on unsupported types."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return load_pdf(file_bytes)
    if ext == "docx":
        return load_docx(file_bytes)
    if ext in {"txt", "md"}:
        return load_txt(file_bytes)
    raise ValueError(
        f"Unsupported file type '.{ext}'. Please upload a .pdf, .docx, .txt or .md file."
    )
